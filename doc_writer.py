import difflib
from typing import List, Tuple, Dict

import docx
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


def get_or_create_hyperlink_style(d):
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = False
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = False
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs
    return "Hyperlink"


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text
    new_run.style = get_or_create_hyperlink_style(part.document)
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def compare2(a, b):
    # list of type: (color: char) if color is black, char is the same in both strings, else it's red
    res = []
    for i, s in enumerate(difflib.ndiff(a, b)):
        if s[0] == ' ':
            res.append(('BLACK', s[-1]))
        else:
            res.append(('RED', s[-1]))
    return res

def try_write_repeats_comparison(document: Document, region_id, repeats_regions: Dict[str, List[str]], b):
    if repeats_regions and region_id in repeats_regions:
        sequences = repeats_regions[region_id]
        if b in sequences:
            return True
        if len(sequences) > 0:
            sequences_str = "\n".join(sequences)
            document.add_paragraph(f"RepeatsDB sequences:\n{sequences_str}")
        for sequence in sequences:
            comp = compare2(sequence, b)
            paragraph = document.add_paragraph('\nComparison with RepeatsDB sequence:\n')
            for color, c in comp:
                run = paragraph.add_run(c)
                if color.upper() == 'RED':
                    run.font.color.rgb = RED
                elif color.upper() == 'BLACK':
                    run.font.color.rgb = BLACK
            document.add_paragraph()
        return False

def write_docx(path, errors: List[Tuple[str, Tuple[str, str]]], missing, df_len: int,
               repeats_regions: Dict[str, List[str]]):
    doc = Document()
    doc.add_heading('RepeatsDB Structure Checker', level=1)
    compare_errors = set()
    errors_that_match_with_repeats = set()
    for region_id, (a, b) in errors:
        h = doc.add_heading(level=3)
        add_hyperlink(h,
                      f"Region ID: {region_id}", f"https://repeatsdb.bio.unipd.it/structure/{region_id.split('_')[0]}"
                      )
        doc.add_paragraph(f"FASTA:\n{a}")
        doc.add_paragraph(f"From query:\n{b}")
        try:
            comparison = compare2(a, b)
        except Exception:
            try:
                if try_write_repeats_comparison(doc, region_id.split('_')[0], repeats_regions, b):
                    errors_that_match_with_repeats.add(region_id)
                    doc.add_paragraph("Sequence matches with RepeatsDB sequence")
            except Exception:
                compare_errors.add(region_id)
            finally:
                continue
        paragraph = doc.add_paragraph('\nComparison with fasta:\n')
        for color, c in comparison:
            run = paragraph.add_run(c)
            if color.upper() == 'RED':
                run.font.color.rgb = RED
            elif color.upper() == 'BLACK':
                run.font.color.rgb = BLACK
        doc.add_paragraph()
        try:
            if try_write_repeats_comparison(doc, region_id.split('_')[0], repeats_regions, b):
                errors_that_match_with_repeats.add(region_id)
                doc.add_paragraph("Sequence matches with RepeatsDB sequence")
        except Exception:
            continue
    if compare_errors:
        doc.add_heading('Comparison errors', level=2)
        for region_id in compare_errors:
            p = doc.add_paragraph()
            add_hyperlink(p, f"COMPARISON ERROR: {region_id}",
                          f"https://repeatsdb.bio.unipd.it/structure/{region_id.split('_')[0]}")
    if missing:
        doc.add_heading('Missing chains', level=2)
        for m in missing:
            p = doc.add_paragraph()
            add_hyperlink(p, f"MISSING: {m}", f"https://repeatsdb.bio.unipd.it/structure/{m}")
    if errors or missing:
        doc.add_heading('Summary', level=2)
        if errors:
            doc.add_paragraph(f"Errors found: {len(errors)}")
            doc.add_paragraph(f"Errors percentage: {len(errors) / df_len * 100}%")
            doc.add_paragraph(f"Errors that match with RepeatsDB: {len(errors_that_match_with_repeats)}")
            doc.add_paragraph(f"Errors percentage that match with RepeatsDB: {len(errors_that_match_with_repeats) / len(errors) * 100}%")
        if missing:
            doc.add_paragraph(f"Missing chains: {len(missing)}")
            doc.add_paragraph(f"Missing chains percentage: {len(missing) / df_len * 100}%")
    doc.save(path)


def load_regions_from_txt(path):
    res = {}
    with open(path, "r") as file:
        lines = file.readlines()
        key = ""
        for line in lines:
            if line.startswith("Region ID:"):
                key = line.split(":")[1].strip()
                res[key] = []
            elif key and line.strip():
                res[key].append(line.strip())
    return res

